import os
from datetime import datetime
from pathlib import Path
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from PyPDF2 import PdfReader
from docx import Document

PASTA_BACKUP = str(Path.home() / "Documents" / "Federal Andaimes - Sistema" / "Backups")
PASTA_BACKUP_CLIENTE = os.path.join(PASTA_BACKUP, "Clientes")
PASTA_BACKUP_PRODUTOS = os.path.join(PASTA_BACKUP, "Produtos")
PASTA_BACKUP_NOTA_FATURA = os.path.join(PASTA_BACKUP, "Notas")

def garantir_pasta_backup():
    try:
        Path(PASTA_BACKUP).mkdir(parents=True, exist_ok=True)
        Path(PASTA_BACKUP_CLIENTE).mkdir(parents=True, exist_ok=True)
        Path(PASTA_BACKUP_PRODUTOS).mkdir(parents=True, exist_ok=True)
        Path(PASTA_BACKUP_NOTA_FATURA).mkdir(parents=True, exist_ok=True)
        return True
    except Exception as e:
        print(f"Erro ao criar pasta de backup: {e}")
        return False

def gerar_nome_arquivo(tipo):
    timestamp = datetime.now().strftime("%d%M%Y_%H%M%S")
    return f"backup_{tipo}_{timestamp}.xlsx"

def exportar_clientes(lista_clientes):
    try:
        if not garantir_pasta_backup():
            return (False, "Erro ao criar pasta de backup")
        if not lista_clientes:
            return (False, "Nenhum cliente cadastrado para exportar")
        nome_arquivo = gerar_nome_arquivo("clientes")
        caminho_completo = os.path.join(PASTA_BACKUP_CLIENTE, nome_arquivo)
        wb = Workbook()
        ws = wb.active
        ws.title = "Clientes"
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_alignment = Alignment(horizontal="center", vertical="center")
        headers = ['ID', 'Nome', 'Email', 'CPF/CNPJ', 'Telefone', 'CEP', 'Endereço', 'Número', 'Complemento', 'Bairro', 'Cidade', 'UF', 'Contrato']
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment
        for row_idx, cliente in enumerate(lista_clientes, start=2):
            for col_idx, valor in enumerate(cliente[:13], start=1):
                ws.cell(row=row_idx, column=col_idx, value=valor)
        column_widths = [8, 30, 30, 18, 18, 12, 35, 10, 20, 20, 20, 5, 20]
        for col_idx, width in enumerate(column_widths, start=1):
            ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width
        wb.save(caminho_completo)
        return (True, caminho_completo)
    except Exception as e:
        return (False, f"Erro ao exportar clientes: {str(e)}")

def exportar_produtos_xls(lista_produtos):
    try:
        if not garantir_pasta_backup():
            return (False, "Erro ao criar pasta de backup")
        if not lista_produtos:
            return (False, "Nenhum produto cadastrado para exportar")
        nome_arquivo = gerar_nome_arquivo("produtos")
        caminho_completo = os.path.join(PASTA_BACKUP_PRODUTOS, nome_arquivo)
        wb = Workbook()
        ws = wb.active
        ws.title = "Produtos"
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_alignment = Alignment(horizontal="center", vertical="center")
        headers = ['ID', 'Nome', 'Preço', 'Unidade', 'Status']
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment
        for row_idx, produto in enumerate(lista_produtos, start=2):
            for col_idx in range(5):
                if col_idx == 4:
                    valor = 'Ativo' if produto[col_idx] == 1 else 'Inativo'
                else:
                    valor = produto[col_idx]
                ws.cell(row=row_idx, column=col_idx + 1, value=valor)
        column_widths = [8, 40, 15, 15, 12]
        for col_idx, width in enumerate(column_widths, start=1):
            ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width
        wb.save(caminho_completo)
        return (True, caminho_completo)
    except Exception as e:
        return (False, f"Erro ao exportar produtos: {str(e)}")
    
def listar_pdfs_nf():
    try:
        pasta_nf = Path.home() / "Documents" / "Federal Andaimes - Sistema" / "Nota Fatura"
        if not pasta_nf.exists():
            return []
        pdfs = []
        for arquivo in pasta_nf.iterdir():
            if arquivo.suffix.lower() == '.pdf':
                pdfs.append(arquivo.name)
        return sorted(pdfs, reverse=True)
    except Exception as e:
        print(f"Erro ao listar PDFs: {e}")
        return []

def converter_pdf_para_docx(nome_arquivo_pdf):
    try:
        pasta_nf = Path.home() / "Documents" / "Federal Andaimes - Sistema" / "Nota Fatura"
        caminho_pdf = pasta_nf / nome_arquivo_pdf
        if not caminho_pdf.exists():
            return (False, "Arquivo PDF não encontrado")
        if not garantir_pasta_backup():
            return (False, "Erro ao criar pasta de backup")
        reader = PdfReader(str(caminho_pdf))
        doc = Document()
        doc.add_heading('Nota Fatura - Conversão de PDF', 0)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                if i > 0:
                    doc.add_page_break()
                doc.add_heading(f'Página {i + 1}', level=1)
                for linha in text.split('\n'):
                    if linha.strip():
                        doc.add_paragraph(linha)
        nome_docx = nome_arquivo_pdf.replace('.pdf', '.docx')
        caminho_docx = os.path.join(PASTA_BACKUP_NOTA_FATURA, nome_docx)
        doc.save(caminho_docx) 
        return (True, caminho_docx) 
    except Exception as e:
        return (False, f"Erro ao converter PDF: {str(e)}")